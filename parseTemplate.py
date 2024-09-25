from studentData import Lecture
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches


class parseTemplate():
    _listOfLecture = []
    _fileName = ""
    def __init__(self,fileName):
        self._fileName = fileName
        self._listOfLecture = []
        self.readListOfLectures()
    def getListOfLectures(self):
        return self._listOfLecture
    def readListOfLectures(self):
        temp = Document(self._fileName)
        tables = temp.tables
    
        paragraphs = temp.paragraphs
        i=0

        #debug printing
        for par in paragraphs:
            print(i, par.text)
            i=i+1

        for i in range(len(tables)-2):
            print("TABLE -----------  ")
            countRow = 0
            for row in tables[i].rows:
                j = 0
                toSkip = False
                for cell in row.cells:
                    print(j, cell.text, end = "|")
                    if "Disciplina promov" in cell.text:
                        toSkip = True
                    if "Disciplina echiva" in cell.text:
                        toSkip = True
                    j = j+1
                print()
                if toSkip:
                    print("SKIPPED")
                    continue
                lectureName = row.cells[6].text
                l = Lecture(lectureName,i+1, "-",row.cells[7].text,row.cells[8].text,"-",row.cells[10].text)
                self._listOfLecture.append(l)
    
    def printTableNotEquivalent(self,doc, notEquivalent):
        table = doc.add_table(1, 7)
        table.style = "TableGrid"
        row = table.rows[0]
        cells = row.cells
        cells[0].text = "Disciplina"
        cells[1].text = "ORE CURS"
        cells[2].text = "ORE L/P/S"
        cells[3].text = "Nota"
        cells[4].text = "Credite"
        cells[5].text = "Anul"
        cells[6].text = "Semestrul"

        for l in notEquivalent:
            row = table.add_row()
            cells = row.cells
            cells[0].text = l._name
            cells[1].text = str(l._nrHLec)
            cells[2].text = l._nrHPrac
            cells[3].text = str(l._grade)
            cells[4].text = str(l._credit)
            cells[5].text = str(l._year)
            cells[6].text = l._sem
            
        pass
    
    def printTableForAYearAndGetCredits(self, table, equivalent):
        countNr= 0 
        credit = 0
        for row in table.rows:
            lectureName = row.cells[6].text
            
            
            found = False
            
            equiFound = ""
            for equi in equivalent:
                if equi[1]._name == lectureName:
                    equiFound = equi
                    found = True
                    break
            if found == False:
                #should not exist
                print("NOT FOUND:",lectureName)
                continue

            cells = row.cells

            countNr = countNr + 1
            cells[0].text = str(countNr)

            l = equiFound[0]
            if l._name == "":
                continue
            cells[1].text = l._name
            cells[2].text = l._nrHLec
            cells[3].text = l._nrHPrac
            cells[4].text = l._grade
            cells[5].text = l._credit
            cells[9].text = l._grade


            try:
                credit = credit + int(equiFound[1]._credit)
            except:
                print("NO CREDIT FOR ",equiFound[1]._name)
            
        return credit

    def writeDocument(self,student, equivalent, notEquivalent):
        doc = Document(self._fileName)


        countYear = 0
        creditList = []
        #update table for each year
        for table in doc.tables:
            if countYear >= student._year:
                continue
            credit = self.printTableForAYearAndGetCredits(table, equivalent)

            creditList.append(credit)
            countYear = countYear+1

        #update table for not equivalent
        tableToUpdate = doc.tables[-2]
        tableToUpdate.style = "TableGrid"
        creditList.append(0)
        countNr = 1 
        for eq in equivalent:
            if eq[0]._id == "":
                row = tableToUpdate.add_row()
                cells = row.cells
                l = eq[1]
                cells[0].text = str(countNr)
                cells[1].text = l._name
                cells[2].text = str(l._year)
                cells[3].text = l._sem
                cells[4].text = l._credit

                countNr = countNr + 1
                try:
                    creditList[-1] = creditList[-1] + int(eq[1]._credit)
                except:
                    print("NO CREDIT FOR ",eq[1]._name)

        self.printTableNotEquivalent(doc, notEquivalent)

        #print text parragraphs
        countYear = 0
        for par in doc.paragraphs:
            text = par.text
            if text.find("Student/ă:") != -1:
                par.text = par.text + student._name
            if text == "Numărul matricol vechi:":
                par.text = par.text + student._nr
            if text.find("Anul de studiu _____")!=-1:
                if (countYear < len(student._years)):
                    par.text = student._years[countYear] +";"+" "*30+ text[text.find(";")+1:]
                    countYear = countYear + 1
            if text.find("Total credite ECTS echivalate:") != -1:
                par.text = text + str(creditList[countYear-1])
            if text.find("Nr. total de credite care urmează a fi obținute") != -1:
                par.text = text[:-1] + ": "+str(creditList[-1])
        
        doc.save("./inputFile/" + student._name[:student._name.find(" ")]+ ".docx")
